"""
Извлечение текста из DOC и DOCX файлов.
- DOCX читается напрямую через python-docx
- DOC конвертируется в DOCX через LibreOffice (если установлен)
"""
from pathlib import Path
import subprocess
import shutil
import tempfile
from docx import Document
from text_normalizer import normalize_legal_text


def extract_docx(file_path: str) -> list[dict]:
    """
    Извлекает текст из DOCX файла.
    Возвращает список "страниц" (условных блоков по ~2000 символов).
    """
    path = Path(file_path)
    source_name = path.name
    
    try:
        doc = Document(str(path))
    except Exception as e:
        print(f"  ⚠️ Ошибка чтения {source_name}: {e}")
        return []
    
    pages = []
    current_text = []
    current_len = 0
    page_num = 1
    TARGET_PAGE_SIZE = 2000  # Условный размер "страницы"
    
    for para in doc.paragraphs:
        text = normalize_legal_text(para.text.strip())
        if not text:
            continue
        
        current_text.append(text)
        current_len += len(text)
        
        if current_len >= TARGET_PAGE_SIZE:
            pages.append({
                "source": source_name,
                "text": normalize_legal_text("\n\n".join(current_text)),
                "page": page_num,
            })
            page_num += 1
            current_text = []
            current_len = 0
    
    if current_text:
        pages.append({
            "source": source_name,
            "text": normalize_legal_text("\n\n".join(current_text)),
            "page": page_num,
        })
    
    # Если документ пустой (только таблицы) — извлекаем из таблиц
    if not pages:
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = normalize_legal_text(cell.text.strip())
                    if cell_text:
                        table_text.append(cell_text)
        
        if table_text:
            pages.append({
            "source": source_name,
            "text": normalize_legal_text("\n\n".join(table_text)),
            "page": 1,
        })
    
    print(f"  📄 {source_name}: извлечено {len(pages)} 'страниц'")
    return pages


def _find_libreoffice() -> str | None:
    """Ищет путь к LibreOffice на Windows/Linux/macOS."""
    # Windows: типичные пути установки
    windows_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for p in windows_paths:
        if Path(p).exists():
            return p
    
    # Пытаемся найти через PATH
    soffice = shutil.which("soffice")
    if soffice:
        return soffice
    
    return None


def extract_doc(file_path: str) -> list[dict]:
    """
    Конвертирует DOC в DOCX через LibreOffice и извлекает текст.
    Если LibreOffice не установлен — пропускает файл с предупреждением.
    """
    path = Path(file_path)
    source_name = path.name
    
    soffice = _find_libreoffice()
    if not soffice:
        print(f"  ⚠️ {source_name}: пропущен (нужен LibreOffice для конвертации .doc)")
        print(f"     Установите LibreOffice: https://www.libreoffice.org/download/")
        print(f"     Или конвертируйте .doc в .docx вручную через MS Word.")
        return []
    
    print(f"  🔄 Конвертация {source_name} (DOC → DOCX)...")
    
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Копируем файл во временную папку
            temp_doc_path = Path(tmp_dir) / path.name
            temp_doc_path.write_bytes(path.read_bytes())
            
            # Конвертируем через LibreOffice в headless режиме
            result = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", tmp_dir,
                    str(temp_doc_path),
                ],
                capture_output=True,
                timeout=60,
                text=True,
            )
            
            if result.returncode != 0:
                print(f"  ❌ Ошибка конвертации: {result.stderr}")
                return []
            
            # Ищем сконвертированный файл
            converted_path = Path(tmp_dir) / (path.stem + ".docx")
            if not converted_path.exists():
                # LibreOffice иногда создает файлы с другим именем
                docx_files = list(Path(tmp_dir).glob("*.docx"))
                if docx_files:
                    converted_path = docx_files[0]
                else:
                    print(f"  ❌ Сконвертированный файл не найден")
                    return []
            
            # Извлекаем текст
            pages = extract_docx(str(converted_path))
            for page in pages:
                page["source"] = source_name
            
            return pages
            
    except subprocess.TimeoutExpired:
        print(f"  ❌ Конвертация {source_name} превысила таймаут")
        return []
    except Exception as e:
        print(f"  ❌ Ошибка при конвертации {source_name}: {e}")
        return []


def extract_all_docs(folder: str) -> list[dict]:
    """
    Извлекает текст из всех DOC и DOCX файлов в указанной папке.
    """
    folder_path = Path(folder)
    if not folder_path.exists():
        print(f"Папка {folder} не найдена.")
        return []
    
    all_pages = []
    doc_files = list(folder_path.glob("*.doc")) + list(folder_path.glob("*.docx"))
    
    if not doc_files:
        print(f"В папке {folder} не найдено DOC/DOCX файлов.")
        return []
    
    print(f"\nНайдено Word-документов: {len(doc_files)}")
    
    # Сначала проверяем, установлен ли LibreOffice (нужен только для .doc)
    doc_files_list = [f for f in doc_files if f.suffix.lower() == ".doc"]
    if doc_files_list:
        if not _find_libreoffice():
            print(f"\n⚠️  Найдено {len(doc_files_list)} файлов .doc")
            print(f"   Для их обработки установите LibreOffice: https://www.libreoffice.org/download/")
            print(f"   Или сконвертируйте их в .docx вручную через MS Word.\n")
    
    for file_path in sorted(doc_files):
        ext = file_path.suffix.lower()
        
        if ext == ".docx":
            pages = extract_docx(str(file_path))
        elif ext == ".doc":
            pages = extract_doc(str(file_path))
        else:
            continue
        
        all_pages.extend(pages)
    
    return all_pages